from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def save_newsletter_to_word(newsletter_text: str, filename="newsletter.docx"):
    doc = Document()

    doc_title = "Newsletter"
    title_paragraph = doc.add_paragraph(doc_title, style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    lines = newsletter_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("**") and line.endswith("**"):
            heading_text = line.strip('*')
            p = doc.add_paragraph(heading_text, style='Heading 1')
            run = p.runs[0]
            run.font.color.rgb = RGBColor(0, 51, 102)
        else:
            p = doc.add_paragraph()
            bold_parts = re.findall(r"\*\*(.*?)\*\*", line)
            if bold_parts:
                p.clear()
                parts = re.split(r"(\*\*.*?\*\*)", line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part.strip('*'))
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.size = Pt(12)
            else:
                run = p.add_run(line)
                run.font.size = Pt(12)

    doc.save(filename)
    print(f"Newsletter saved to {filename}")
